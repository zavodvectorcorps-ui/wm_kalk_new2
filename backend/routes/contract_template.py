"""Contract template management - upload, placeholder extraction, mapping settings, generation with KP attachment."""
from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime, timezone
from database import db
import os
import io
import re
import logging
import httpx

router = APIRouter(prefix="/sauna-crm/contract-template", tags=["Contract Template"])
logger = logging.getLogger(__name__)

TEMPLATE_DIR = "/app/backend/templates"
TEMPLATE_FILENAME = "contract_template.docx"

# All available lead fields for mapping
AVAILABLE_SOURCES = [
    {"id": "clientName", "label": "Имя клиента", "category": "client"},
    {"id": "phone", "label": "Телефон", "category": "client"},
    {"id": "email", "label": "Email", "category": "client"},
    {"id": "address", "label": "Адрес", "category": "client"},
    {"id": "modelName", "label": "Модель", "category": "production"},
    {"id": "totalAmount", "label": "Сумма заказа", "category": "payment"},
    {"id": "advancePayment", "label": "Предоплата", "category": "payment"},
    {"id": "paidAmount", "label": "Оплачено", "category": "payment"},
    {"id": "prepaymentDate", "label": "Дата предоплаты", "category": "payment"},
    {"id": "readyDate", "label": "Дата готовности", "category": "production"},
    {"id": "productionDate", "label": "Дата производства", "category": "production"},
    {"id": "deliveryDate", "label": "Дата доставки", "category": "production"},
    {"id": "notes", "label": "Заметки", "category": "other"},
    {"id": "amocrm_id", "label": "amoCRM ID", "category": "other"},
    {"id": "id", "label": "ID лида", "category": "other"},
    {"id": "field_1", "label": "Поле 1 (CRM)", "category": "crm_fields"},
    {"id": "field_2", "label": "Поле 2 (CRM)", "category": "crm_fields"},
    {"id": "field_3", "label": "Поле 3 (CRM)", "category": "crm_fields"},
    {"id": "field_4", "label": "Поле 4 (CRM)", "category": "crm_fields"},
    {"id": "field_5", "label": "Поле 5 (CRM)", "category": "crm_fields"},
    {"id": "field_6", "label": "Поле 6 (CRM)", "category": "crm_fields"},
    {"id": "field_7", "label": "Поле 7 (CRM)", "category": "crm_fields"},
    {"id": "field_8", "label": "Поле 8 (CRM)", "category": "crm_fields"},
    {"id": "field_9", "label": "Поле 9 (CRM)", "category": "crm_fields"},
    {"id": "field_10", "label": "Поле 10 (CRM)", "category": "crm_fields"},
    # Computed / calculator fields
    {"id": "_contract_date", "label": "Дата договора (сегодня)", "category": "computed"},
    {"id": "_deposit_percent", "label": "Процент задатка", "category": "computed"},
    {"id": "_offer_number", "label": "Номер предложения", "category": "computed"},
    {"id": "_calc_width", "label": "Ширина (из калькулятора)", "category": "calculator"},
    {"id": "_calc_length", "label": "Длина (из калькулятора)", "category": "calculator"},
    {"id": "_calc_version", "label": "Версия (из калькулятора)", "category": "calculator"},
    {"id": "_calc_model", "label": "Модель (из калькулятора)", "category": "calculator"},
    {"id": "_calc_total_price", "label": "Цена (из калькулятора)", "category": "calculator"},
]

SETTINGS_COLLECTION = "contract_template_settings"


def _extract_placeholders_from_docx(filepath: str) -> list:
    """Extract all {{...}} placeholders from a DOCX file, handling split runs."""
    from docx import Document
    doc = Document(filepath)
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
    # Also check headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    all_text += para.text + "\n"
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    all_text += para.text + "\n"
    placeholders = sorted(set(re.findall(r'\{\{[A-Z0-9_]+\}\}', all_text)))
    return placeholders


DEFAULT_MAPPINGS = {
    "{{CONTRACT_DATE}}": {"source": "_contract_date", "defaultValue": "", "label": "Дата договора"},
    "{{CONTRACT_CITY}}": {"source": "_static", "defaultValue": "Warszawie", "label": "Город"},
    "{{CLIENT_NAME}}": {"source": "clientName", "defaultValue": "...............", "label": "Имя клиента"},
    "{{CLIENT_ADDRESS}}": {"source": "address", "defaultValue": "...............", "label": "Адрес клиента"},
    "{{SAUNA_TYPE}}": {"source": "modelName", "defaultValue": "...............", "label": "Тип/модель сауны"},
    "{{SAUNA_WIDTH}}": {"source": "_calc_width", "defaultValue": "...", "label": "Ширина"},
    "{{SAUNA_LENGTH}}": {"source": "_calc_length", "defaultValue": "...", "label": "Длина"},
    "{{SAUNA_VERSION}}": {"source": "_calc_version", "defaultValue": "[wersja gotowa zlozona]", "label": "Версия"},
    "{{OFFER_NUMBER}}": {"source": "_offer_number", "defaultValue": "...............", "label": "Номер предложения"},
    "{{TOTAL_PRICE}}": {"source": "totalAmount", "defaultValue": "0", "label": "Общая цена"},
    "{{DEPOSIT_PERCENT}}": {"source": "_deposit_percent", "defaultValue": "30", "label": "Процент задатка"},
    "{{DEPOSIT_AMOUNT}}": {"source": "advancePayment", "defaultValue": "0", "label": "Сумма задатка"},
    "{{DELIVERY_PAYER}}": {"source": "_static", "defaultValue": "Sprzedawcy", "label": "Плательщик доставки"},
}


async def _get_settings() -> dict:
    settings = await db[SETTINGS_COLLECTION].find_one({"type": "contract_template"}, {"_id": 0})

    # Get template placeholders
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    placeholders = []
    if os.path.exists(template_path):
        placeholders = _extract_placeholders_from_docx(template_path)

    if not settings:
        # First time — create defaults
        mappings = []
        for ph in placeholders:
            if ph in DEFAULT_MAPPINGS:
                m = DEFAULT_MAPPINGS[ph]
                mappings.append({"placeholder": ph, **m})
            else:
                mappings.append({
                    "placeholder": ph, "source": "_static",
                    "defaultValue": "", "label": ph.strip("{}")
                })

        settings = {
            "type": "contract_template",
            "templateName": TEMPLATE_FILENAME,
            "uploadedAt": None,
            "mappings": mappings,
            "attachKp": True,
            "placeholders": placeholders
        }
        await db[SETTINGS_COLLECTION].insert_one({**settings})
    else:
        # Existing settings — ensure ALL template placeholders are covered
        existing_phs = {m["placeholder"] for m in settings.get("mappings", [])}
        missing = [ph for ph in placeholders if ph not in existing_phs]
        if missing:
            new_mappings = list(settings.get("mappings", []))
            for ph in missing:
                if ph in DEFAULT_MAPPINGS:
                    new_mappings.append({"placeholder": ph, **DEFAULT_MAPPINGS[ph]})
                else:
                    new_mappings.append({
                        "placeholder": ph, "source": "_static",
                        "defaultValue": "", "label": ph.strip("{}")
                    })
            settings["mappings"] = new_mappings
            await db[SETTINGS_COLLECTION].update_one(
                {"type": "contract_template"},
                {"$set": {"mappings": new_mappings, "placeholders": placeholders}}
            )
        settings["placeholders"] = placeholders

    return settings


@router.get("/settings")
async def get_template_settings():
    """Get contract template settings including field mappings."""
    settings = await _get_settings()
    return {
        **settings,
        "availableSources": AVAILABLE_SOURCES
    }


@router.post("/settings")
async def save_template_settings(request: dict):
    """Save contract template settings (mappings, attachKp flag)."""
    mappings = request.get("mappings", [])
    attach_kp = request.get("attachKp", True)

    await db[SETTINGS_COLLECTION].update_one(
        {"type": "contract_template"},
        {"$set": {
            "mappings": mappings,
            "attachKp": attach_kp,
            "updatedAt": datetime.now(timezone.utc).isoformat()
        }},
        upsert=True
    )
    return {"status": "ok"}


@router.post("/upload")
async def upload_template(file: UploadFile = File(...)):
    """Upload a new DOCX contract template. Stores in Cloudinary for persistence."""
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are allowed")

    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    # Save locally (for immediate use)
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    with open(template_path, "wb") as f:
        f.write(content)

    # Upload to Cloudinary for persistence across deploys
    template_url = None
    from services.cloudinary_service import is_cloudinary_configured
    if is_cloudinary_configured():
        try:
            import cloudinary.uploader
            result = cloudinary.uploader.upload(
                io.BytesIO(content),
                resource_type="raw",
                folder="contract_templates",
                public_id=f"contract_template_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}",
                format="docx"
            )
            template_url = result.get("secure_url")
            logger.info(f"Template uploaded to Cloudinary: {template_url}")
        except Exception as e:
            logger.error(f"Cloudinary upload failed: {e}")

    # Extract placeholders
    placeholders = _extract_placeholders_from_docx(template_path)

    # Update settings — keep existing mappings, add new placeholders
    existing = await _get_settings()
    existing_map = {m["placeholder"]: m for m in existing.get("mappings", [])}

    new_mappings = []
    for ph in placeholders:
        if ph in existing_map:
            new_mappings.append(existing_map[ph])
        elif ph in DEFAULT_MAPPINGS:
            new_mappings.append({"placeholder": ph, **DEFAULT_MAPPINGS[ph]})
        else:
            new_mappings.append({
                "placeholder": ph,
                "source": "_static",
                "defaultValue": "",
                "label": ph.replace("{{", "").replace("}}", "")
            })

    await db[SETTINGS_COLLECTION].update_one(
        {"type": "contract_template"},
        {"$set": {
            "templateName": file.filename,
            "templateUrl": template_url,
            "uploadedAt": datetime.now(timezone.utc).isoformat(),
            "mappings": new_mappings,
            "placeholders": placeholders
        }},
        upsert=True
    )

    return {
        "status": "ok",
        "templateName": file.filename,
        "templateUrl": template_url,
        "placeholders": placeholders,
        "mappingsCount": len(new_mappings)
    }


@router.get("/placeholders")
async def get_placeholders():
    """Extract placeholders from current template."""
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail="Template not found")
    placeholders = _extract_placeholders_from_docx(template_path)
    return {"placeholders": placeholders}



@router.get("/debug/{lead_id}")
async def debug_contract_generation(lead_id: str):
    """Diagnostic endpoint to check contract generation prerequisites for a specific lead."""
    import traceback
    try:
        result = {"lead_id": lead_id, "checks": {}}

        # 1. Check lead exists
        lead = await db.sauna_crm_leads.find_one({"id": lead_id}, {"_id": 0})
        if not lead:
            result["checks"]["lead"] = {"status": "ERROR", "detail": "Lead not found"}
            return result

        # List all document types in the lead
        doc_list = []
        for d in lead.get("documents", []):
            doc_list.append({
                "type": d.get("type"),
                "name": d.get("name"),
                "url": str(d.get("url", ""))[:120]
            })

        result["checks"]["lead"] = {
            "status": "OK",
            "clientName": lead.get("clientName"),
            "documents": doc_list,
            "calculatorPdfUrl": lead.get("calculatorPdfUrl"),
            "calculatorOrderId": lead.get("calculatorOrderId"),
            "calculatorCollection": lead.get("calculatorCollection"),
        }

        # 2. Check template
        try:
            settings = await _get_settings()
            template_url = settings.get("templateUrl")
            template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
            local_exists = os.path.exists(template_path)
            result["checks"]["template"] = {
                "status": "OK" if (template_url or local_exists) else "ERROR",
                "templateUrl": template_url,
                "localFileExists": local_exists,
                "mappings_count": len(settings.get("mappings", [])),
                "attachKp": settings.get("attachKp", True),
            }
        except Exception as e:
            result["checks"]["template"] = {"status": "ERROR", "error": str(e)}

        # 3. Check KP sources
        kp_url = None
        kp_source = "none"
        kp_types = ("kp", "commercial_proposal", "кп", "kp_pdf")
        for doc_item in lead.get("documents", []):
            doc_type = (doc_item.get("type") or "").lower()
            doc_name = (doc_item.get("name") or "").lower()
            if doc_type in kp_types or "кп" in doc_name or "kp" in doc_name:
                kp_url = doc_item.get("url")
                kp_source = f"documents (type={doc_item.get('type')}, name={doc_item.get('name')})"
                break
        if not kp_url:
            kp_url = lead.get("calculatorPdfUrl")
            if kp_url:
                kp_source = "calculatorPdfUrl"

        is_proxy = bool(kp_url and ("calculator-pdf/" in kp_url or kp_url.startswith("/api/")))
        result["checks"]["kp"] = {
            "kp_url": kp_url,
            "source": kp_source,
            "is_proxy_url": is_proxy,
        }

        # 4. Check calculator_pdfs collection
        calc_order_id = lead.get("calculatorOrderId")
        if calc_order_id:
            try:
                pdf_doc = await db["calculator_pdfs"].find_one(
                    {"order_id": calc_order_id},
                    {"_id": 0, "order_id": 1, "cloudinary_url": 1, "created_at": 1}
                )
                has_pdf_data = False
                if pdf_doc is not None:
                    pdf_doc_check = await db["calculator_pdfs"].find_one(
                        {"order_id": calc_order_id},
                        {"_id": 0, "pdf_data": 1}
                    )
                    has_pdf_data = bool(pdf_doc_check and pdf_doc_check.get("pdf_data"))

                # Also try to extract order_id from proxy URL if different
                proxy_order_id = None
                if kp_url and is_proxy:
                    proxy_order_id = kp_url.rstrip("/").split("/")[-1]

                result["checks"]["calculator_pdfs"] = {
                    "calc_order_id": calc_order_id,
                    "proxy_order_id": proxy_order_id,
                    "found_by_calc_id": pdf_doc is not None,
                    "has_pdf_data": has_pdf_data,
                    "cloudinary_url": (pdf_doc or {}).get("cloudinary_url"),
                }

                # Check if proxy order ID is different and has data
                if proxy_order_id and proxy_order_id != calc_order_id:
                    pdf_doc2 = await db["calculator_pdfs"].find_one(
                        {"order_id": proxy_order_id},
                        {"_id": 0, "order_id": 1, "cloudinary_url": 1}
                    )
                    result["checks"]["calculator_pdfs"]["found_by_proxy_id"] = pdf_doc2 is not None
            except Exception as e:
                result["checks"]["calculator_pdfs"] = {"status": "ERROR", "error": str(e)}
        else:
            result["checks"]["calculator_pdfs"] = {"status": "SKIP", "reason": "No calculatorOrderId"}

        # 5. Check Cloudinary
        from services.cloudinary_service import is_cloudinary_configured
        result["checks"]["cloudinary"] = {"configured": is_cloudinary_configured()}

        return result
    except Exception as e:
        return {"error": str(e), "traceback": traceback.format_exc()}



def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholders in a paragraph, handling split runs."""
    full_text = paragraph.text
    has_placeholder = False
    for ph in replacements:
        if ph in full_text:
            has_placeholder = True
            break
    if not has_placeholder:
        return

    # Try direct run replacement first
    for run in paragraph.runs:
        for ph, val in replacements.items():
            if ph in run.text:
                run.text = run.text.replace(ph, val)

    # Check if any placeholders still remain (split across runs)
    remaining_text = paragraph.text
    needs_merge = False
    for ph in replacements:
        if ph in full_text and ph in remaining_text:
            # Already replaced in runs above — skip
            pass
        elif ph in full_text and ph not in remaining_text:
            # Successfully replaced
            pass
        elif "{{" in remaining_text:
            needs_merge = True
            break

    if needs_merge and paragraph.runs:
        # Rebuild: merge all runs, do replacement, put text in first run
        merged = "".join(run.text for run in paragraph.runs)
        for ph, val in replacements.items():
            merged = merged.replace(ph, val)
        # Keep first run's formatting, clear others
        paragraph.runs[0].text = merged
        for run in paragraph.runs[1:]:
            run.text = ""


def _resolve_value(source: str, default_value: str, lead: dict, calc_order: dict | None, now: datetime) -> str:
    """Resolve a mapping source to an actual value."""
    if source == "_static":
        return default_value

    if source == "_contract_date":
        return now.strftime("%d.%m.%Y")

    if source == "_deposit_percent":
        total = float(lead.get("totalAmount") or (calc_order or {}).get("totalPrice") or 0)
        advance = float(lead.get("advancePayment") or lead.get("prepayment") or 0)
        if total and advance:
            try:
                return str(round(advance / total * 100))
            except (ValueError, ZeroDivisionError):
                pass
        return default_value or "30"

    if source == "_offer_number":
        return lead.get("calculatorOrderId") or lead.get("id") or default_value

    if source == "_calc_width":
        if calc_order:
            return str(calc_order.get("width", calc_order.get("szerokosc", ""))) or default_value
        return default_value

    if source == "_calc_length":
        if calc_order:
            return str(calc_order.get("length", calc_order.get("dlugosc", ""))) or default_value
        return default_value

    if source == "_calc_version":
        if calc_order:
            return str(calc_order.get("version", calc_order.get("wersja", ""))) or default_value
        return default_value

    if source == "_calc_model":
        if calc_order:
            return calc_order.get("model", calc_order.get("modelName", "")) or default_value
        return default_value

    if source == "_calc_total_price":
        if calc_order:
            val = calc_order.get("totalPrice", calc_order.get("total_price", 0))
            return _fmt_amount(val) if val else default_value
        return default_value

    # Direct lead field
    val = lead.get(source)
    if val is None or val == "":
        return default_value or ""

    # Format amounts for money fields
    if source in ("totalAmount", "advancePayment", "paidAmount"):
        return _fmt_amount(val)

    return str(val)


def _fmt_amount(val) -> str:
    try:
        v = float(val)
        if v == int(v):
            return f"{int(v):,}".replace(",", " ")
        return f"{v:,.2f}".replace(",", " ")
    except (ValueError, TypeError):
        return str(val)


async def _download_file(url: str) -> bytes | None:
    """Download a file from URL with retry."""
    for attempt in range(2):
        try:
            async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
                resp = await client.get(url)
                logger.info(f"Download {url}: status={resp.status_code}, size={len(resp.content)}")
                if resp.status_code == 200 and len(resp.content) > 0:
                    return resp.content
                logger.warning(f"Download failed: status={resp.status_code}")
        except Exception as e:
            logger.error(f"Download attempt {attempt+1} failed for {url}: {e}")
    return None


def _pdf_to_images(pdf_bytes: bytes) -> list:
    """Convert PDF pages to PNG images using PyMuPDF. Returns list of (bytes, width, height)."""
    import fitz
    images = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render at 2x resolution for quality
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            images.append((img_bytes, pix.width, pix.height))
        doc.close()
    except Exception as e:
        logger.error(f"PDF to images conversion failed: {e}")
    return images


async def generate_contract_with_kp(lead_id: str) -> dict:
    """Generate contract DOCX with dynamic mappings and attached KP PDF pages."""
    import traceback
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from services.cloudinary_service import is_cloudinary_configured

    logger.info(f"=== CONTRACT GENERATION START for lead_id={lead_id} ===")

    lead = await db.sauna_crm_leads.find_one({"id": lead_id}, {"_id": 0})
    if not lead:
        raise HTTPException(status_code=404, detail="Lead not found")

    logger.info(f"Lead found: clientName={lead.get('clientName')}, docs={len(lead.get('documents', []))}")

    # Get calculator order
    calc_order = None
    calc_order_id = lead.get("calculatorOrderId")
    calc_col = lead.get("calculatorCollection", "sauna_orders")
    if calc_order_id:
        calc_order = await db[calc_col].find_one({"id": calc_order_id}, {"_id": 0})
        logger.info(f"Calculator order: {'found' if calc_order else 'NOT found'} (id={calc_order_id}, col={calc_col})")

    # Load template — prefer user-uploaded (from Cloudinary), fallback to local
    settings = await _get_settings()
    mappings = settings.get("mappings", [])
    attach_kp = settings.get("attachKp", True)
    template_url = settings.get("templateUrl")
    logger.info(f"Settings: mappings={len(mappings)}, attachKp={attach_kp}, templateUrl={'yes' if template_url else 'no'}")

    template_data = None
    if template_url:
        logger.info(f"Downloading custom template from: {template_url}")
        try:
            template_data = await _download_file(template_url)
            if template_data:
                logger.info(f"Custom template loaded: {len(template_data)} bytes")
            else:
                logger.warning("Failed to download custom template, falling back to local")
        except Exception as e:
            logger.error(f"Template download error: {e}")

    if template_data:
        doc = Document(io.BytesIO(template_data))
    else:
        template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="Contract template not found on server. Please upload a template first.")
        doc = Document(template_path)
        logger.info(f"Using local template: {template_path}")

    now = datetime.now(timezone.utc)

    # Build replacements from mappings
    replacements = {}
    for m in mappings:
        ph = m["placeholder"]
        source = m.get("source", "_static")
        default_val = m.get("defaultValue", "")
        try:
            replacements[ph] = _resolve_value(source, default_val, lead, calc_order, now)
        except Exception as e:
            logger.error(f"Error resolving placeholder {ph} (source={source}): {e}")
            replacements[ph] = default_val or ""

    logger.info(f"Replacements built: {len(replacements)} placeholders")

    # Replace in paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, replacements)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, replacements)

    logger.info("Placeholder replacement complete")

    # Remove old embedded images (KP pages) if any
    try:
        _remove_trailing_images(doc)
    except Exception as e:
        logger.error(f"Error removing trailing images: {e}")

    # Attach KP PDF as images
    kp_url = None
    kp_attached = False
    kp_error = None
    if attach_kp:
        try:
            kp_url, kp_attached = await _attach_kp_to_doc(doc, lead, calc_order_id)
            if not kp_attached:
                kp_error = f"KP not attached: url={kp_url}, no PDF data available"
        except Exception as e:
            kp_error = f"KP attachment exception: {str(e)}"
            logger.error(f"KP attachment failed (non-fatal): {e}\n{traceback.format_exc()}")
    else:
        kp_error = "attachKp is disabled in settings"

    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    logger.info(f"Document saved to buffer: {docx_buffer.getbuffer().nbytes} bytes")

    # Upload
    file_url = None
    client_name = lead.get("clientName", "")
    if is_cloudinary_configured():
        try:
            import cloudinary.uploader
            result = cloudinary.uploader.upload(
                docx_buffer,
                resource_type="raw",
                folder="contracts",
                public_id=f"contract_{lead_id}_{now.strftime('%Y%m%d_%H%M%S')}",
                format="docx"
            )
            file_url = result.get("secure_url")
            logger.info(f"Contract uploaded to Cloudinary: {file_url}")
        except Exception as e:
            logger.error(f"Cloudinary upload failed: {e}")

    if not file_url:
        local_dir = "/app/backend/static/contracts"
        os.makedirs(local_dir, exist_ok=True)
        fname = f"contract_{lead_id}_{now.strftime('%Y%m%d_%H%M%S')}.docx"
        local_path = os.path.join(local_dir, fname)
        docx_buffer.seek(0)
        with open(local_path, "wb") as f:
            f.write(docx_buffer.read())
        file_url = f"/api/static/contracts/{fname}"
        logger.info(f"Contract saved locally: {file_url}")

    # Update lead documents
    docs = lead.get("documents", [])
    docs = [d for d in docs if d.get("type") != "contract"]
    contract_doc = {
        "type": "contract",
        "name": f"Umowa {client_name}".strip(),
        "url": file_url,
        "createdAt": now.isoformat(),
        "format": "docx"
    }
    docs.append(contract_doc)

    await db.sauna_crm_leads.update_one(
        {"id": lead_id},
        {"$set": {
            "documents": docs,
            "contractUrl": file_url,
            "updatedAt": now.isoformat()
        }}
    )

    logger.info(f"=== CONTRACT GENERATION COMPLETE for lead_id={lead_id}, url={file_url}, kpAttached={kp_attached} ===")

    return {
        "status": "ok",
        "contractUrl": file_url,
        "kpUrl": kp_url,
        "kpAttached": kp_attached,
        "kpError": kp_error,
        "replacements": replacements
    }



async def _attach_kp_to_doc(doc, lead: dict, calc_order_id: str | None) -> tuple:
    """Attach KP PDF pages as images to the contract document. Returns (kp_url, was_attached)."""
    from docx import Document
    from docx.shared import Inches, Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    kp_url = None

    logger.info(f"Looking for KP in lead documents ({len(lead.get('documents', []))} docs)")

    # Find KP from lead documents — check multiple type names
    kp_types = ("kp", "commercial_proposal", "кп", "kp_pdf")
    for doc_item in lead.get("documents", []):
        doc_type = (doc_item.get("type") or "").lower()
        doc_name = (doc_item.get("name") or "").lower()
        if doc_type in kp_types or "кп" in doc_name or "kp" in doc_name:
            kp_url = doc_item.get("url")
            logger.info(f"Found KP document: type={doc_item.get('type')}, url={kp_url}")
            break

    # Fallback to calculator PDF
    if not kp_url:
        kp_url = lead.get("calculatorPdfUrl")
        if kp_url:
            logger.info(f"Using calculatorPdfUrl: {kp_url}")

    if not kp_url:
        logger.info("No KP URL found in lead documents or calculatorPdfUrl")
        return None, False

    pdf_bytes = None

    # Detect proxy URLs (local API endpoints or full URLs containing calculator-pdf/)
    is_proxy = "calculator-pdf/" in kp_url or kp_url.startswith("/api/")

    if not is_proxy and kp_url.startswith("http"):
        # Direct URL (e.g. Cloudinary) — download immediately, no MongoDB needed
        logger.info(f"KP URL is direct (Cloudinary/HTTP), downloading: {kp_url[:100]}")
        pdf_bytes = await _download_file(kp_url)
        if pdf_bytes:
            logger.info(f"Downloaded KP: {len(pdf_bytes)} bytes")
        else:
            logger.warning(f"Failed to download KP from direct URL: {kp_url[:100]}")
    elif is_proxy:
        # Proxy URL — get PDF data from MongoDB to avoid self-referential HTTP calls
        proxy_order_id = kp_url.rstrip("/").split("/")[-1] if "/" in kp_url else calc_order_id
        logger.info(f"KP URL is proxy, reading PDF from MongoDB: order_id={proxy_order_id}")

        if proxy_order_id:
            try:
                pdf_doc = await db["calculator_pdfs"].find_one(
                    {"order_id": proxy_order_id},
                    {"pdf_data": 1, "cloudinary_url": 1}
                )
                if pdf_doc:
                    if pdf_doc.get("cloudinary_url"):
                        kp_url = pdf_doc["cloudinary_url"]
                        logger.info(f"Resolved proxy URL to Cloudinary: {kp_url}")
                    if pdf_doc.get("pdf_data"):
                        pdf_bytes = pdf_doc["pdf_data"]
                        if not isinstance(pdf_bytes, bytes):
                            pdf_bytes = bytes(pdf_bytes)
                        logger.info(f"Got PDF from MongoDB: {len(pdf_bytes)} bytes")
                else:
                    logger.warning(f"PDF not found in calculator_pdfs for order_id={proxy_order_id}")
            except Exception as e:
                logger.error(f"MongoDB query failed for calculator_pdfs: {e}")

        # If MongoDB failed/empty but we resolved to a Cloudinary URL, download it
        if not pdf_bytes and kp_url and kp_url.startswith("http") and "calculator-pdf/" not in kp_url:
            logger.info(f"Falling back to HTTP download of resolved URL: {kp_url[:100]}")
            pdf_bytes = await _download_file(kp_url)

        # Last resort: try by calc_order_id
        if not pdf_bytes and calc_order_id and calc_order_id != proxy_order_id:
            try:
                pdf_doc = await db["calculator_pdfs"].find_one(
                    {"order_id": calc_order_id},
                    {"pdf_data": 1, "cloudinary_url": 1}
                )
                if pdf_doc:
                    if pdf_doc.get("pdf_data"):
                        pdf_bytes = pdf_doc["pdf_data"]
                        if not isinstance(pdf_bytes, bytes):
                            pdf_bytes = bytes(pdf_bytes)
                        logger.info(f"Got PDF from calculator_pdfs by calc_order_id: {len(pdf_bytes)} bytes")
                    if not pdf_bytes and pdf_doc.get("cloudinary_url"):
                        kp_url = pdf_doc["cloudinary_url"]
                        pdf_bytes = await _download_file(kp_url)
            except Exception as e:
                logger.error(f"MongoDB fallback query failed: {e}")

    if not pdf_bytes:
        logger.warning(f"No PDF data obtained for KP (url={kp_url})")
        return kp_url, False

    logger.info(f"Processing KP: {len(pdf_bytes)} bytes, first4={pdf_bytes[:4]}")
    kp_images = []

    if pdf_bytes[:4] == b'%PDF':
        kp_images = _pdf_to_images(pdf_bytes)
        logger.info(f"PDF converted to {len(kp_images)} images")
    elif pdf_bytes[:8] == b'\x89PNG\r\n\x1a\n' or pdf_bytes[:2] == b'\xff\xd8':
        from PIL import Image
        img = Image.open(io.BytesIO(pdf_bytes))
        kp_images = [(pdf_bytes, img.width, img.height)]
        logger.info(f"KP is image: {img.width}x{img.height}")
    else:
        logger.info("Unknown format, trying as PDF")
        kp_images = _pdf_to_images(pdf_bytes)

    if not kp_images:
        logger.warning("No images extracted from KP")
        return kp_url, False

    # Add page break and header
    bp = doc.add_paragraph()
    run = bp.add_run()
    run.add_break(WD_BREAK.PAGE)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("Załącznik nr 1 – Specyfikacja")
    header_run.bold = True
    header_run.font.size = Pt(14)

    for img_bytes_data, w, h in kp_images:
        img_stream = io.BytesIO(img_bytes_data)
        max_width = Inches(6.5)
        aspect = h / w if w > 0 else 1
        img_width = max_width
        img_height = Emu(int(img_width * aspect))
        if img_height > Inches(9):
            img_height = Inches(9)
            img_width = Emu(int(img_height / aspect)) if aspect > 0 else max_width
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img_stream, width=img_width)

    logger.info(f"Added {len(kp_images)} KP pages to contract")
    return kp_url, True


def _remove_trailing_images(doc):
    """Remove image-only paragraphs at the end of the document (old KP images)."""
    from docx.oxml.ns import qn
    # Find the last paragraph that contains "Załącznik" or "ZAŁĄCZNIK"
    zalacznik_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "załącznik" in para.text.lower() or "zalacznik" in para.text.lower():
            zalacznik_idx = i

    if zalacznik_idx < 0:
        logger.info("No 'załącznik' found in document - skipping image removal")
        return

    logger.info(f"Found 'załącznik' at paragraph {zalacznik_idx}: '{doc.paragraphs[zalacznik_idx].text[:50]}'")

    # Remove all paragraphs after the załącznik that contain only images or are empty
    body = doc.element.body
    paras_to_remove = []
    for i in range(len(doc.paragraphs) - 1, zalacznik_idx, -1):
        para = doc.paragraphs[i]
        has_image = bool(para._element.findall(f'.//{qn("wp:inline")}') or para._element.findall(f'.//{qn("wp:anchor")}'))
        is_empty = not para.text.strip()
        if has_image or is_empty:
            paras_to_remove.append(para._element)
        else:
            break

    logger.info(f"Removing {len(paras_to_remove)} trailing image/empty paragraphs")
    for elem in paras_to_remove:
        body.remove(elem)

    # Also remove the "Załącznik" paragraph itself
    if zalacznik_idx >= 0 and zalacznik_idx < len(doc.paragraphs):
        para = doc.paragraphs[zalacznik_idx]
        if "załącznik" in para.text.lower() or "zalacznik" in para.text.lower():
            logger.info(f"Removing 'załącznik' paragraph: '{para.text[:50]}'")
            body.remove(para._element)
